import os
import io
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from io import BytesIO
from fpdf import FPDF
from docx import Document
import pandas as pd
import random

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

@app.route('/upload', methods=['POST'])
def upload_file():
    """Upload and process the Excel file."""
    try:
        file = request.files['file']
        filepath = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(filepath)
        df = pd.read_excel(filepath)

        # Validate columns
        required_columns = {'Department', 'Class', 'Teacher', 'Subject'}
        if not required_columns.issubset(df.columns):
            return jsonify({'error': 'Required columns missing'}), 400

        departments = df['Department'].dropna().unique().tolist()
        classes = df['Class'].dropna().unique().tolist()

        # Cache file
        app.config['uploaded_file'] = filepath
        return jsonify({'departments': departments, 'classes': classes})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/generate_schedule', methods=['POST'])
def generate_schedule():
    """Generate a timetable with teacher and slot constraints."""
    try:
        department = request.json['department']
        selected_class = request.json['class']
        filepath = app.config.get('uploaded_file')

        df = pd.read_excel(filepath)
        data = df[(df['Department'] == department) & (df['Class'] == selected_class)]

        if data.empty:
            return jsonify({'error': 'No data for selected class/department'}), 404

        # Initialize schedule grid
        days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday"]
        times = ["9:00-10:00", "10:00-11:00", "11:00-12:00", "1:00-2:00", "2:00-3:00", "3:00-4:00"]
        timetable = {day: {time: "" for time in times} for day in days}

        # Track teacher assignments to prevent scheduling conflicts across departments and classes
        teacher_assignments = {}

        subjects = data[['Teacher', 'Subject']].drop_duplicates().values.tolist()

        # Function to check for conflicts
        def check_teacher_conflict(teacher, day, time):
            # Check if teacher is already assigned to this time in any class or department
            if (teacher, day, time) in teacher_assignments:
                return True
            return False

        # Assign subjects to time slots, ensuring no conflicts
        for day in days:
            for time in times:
                random.shuffle(subjects)
                for teacher, subject in subjects:
                    # Check if teacher is already assigned to this time in any class/department
                    if not check_teacher_conflict(teacher, day, time):
                        timetable[day][time] = f"{subject} ({teacher})"
                        # Assign the teacher to this time slot globally across all departments and classes
                        teacher_assignments[(teacher, day, time)] = selected_class
                        break

        # Format for frontend
        slots = [{"time": time, "schedule": [timetable[day][time] for day in days]} for time in times]
        return jsonify({"days": days, "slots": slots})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download', methods=['POST'])
def download_file():
    """Export the provided schedule to a file format."""
    try:
        # Parse request data
        format = request.json['format']
        schedule = request.json['schedule']

        if not schedule:
            return jsonify({'error': 'No schedule data provided'}), 400

        # Extract schedule data
        days = schedule['days']
        slots = schedule['slots']

        # Convert schedule to DataFrame
        schedule_df = pd.DataFrame([
            {**{"Time": slot["time"]}, **{day: item for day, item in zip(days, slot["schedule"])}}
            for slot in slots
        ])

        # Prepare the file for the requested format
        buffer = BytesIO()

        if format == "csv":
            schedule_df.to_csv(buffer, index=False)
            buffer.seek(0)
            return send_file(buffer, download_name="schedule.csv", as_attachment=True)

        elif format == "word":
            doc = Document()
            doc.add_heading("Generated Schedule", level=1)
            table = doc.add_table(rows=schedule_df.shape[0] + 1, cols=schedule_df.shape[1])
            
            # Add headers
            for i, column in enumerate(schedule_df.columns):
                table.cell(0, i).text = column
            
            # Add rows of data
            for i, row in schedule_df.iterrows():
                for j, value in enumerate(row):
                    table.cell(i + 1, j).text = str(value)
            
            doc.save(buffer)
            buffer.seek(0)
            return send_file(buffer, download_name="schedule.docx", as_attachment=True)

        elif format == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)

            # Add headers
            for col in schedule_df.columns:
                pdf.cell(40, 10, col, 1)
            pdf.ln()

            # Add rows of data
            for _, row in schedule_df.iterrows():
                for val in row:
                    pdf.cell(40, 10, str(val), 1)
                pdf.ln()

            pdf.output(buffer)
            buffer.seek(0)
            return send_file(buffer, download_name="schedule.pdf", as_attachment=True)

        else:
            return jsonify({"error": "Invalid format requested"}), 400

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download_entire_timetable', methods=['POST'])
def download_entire_timetable():
    data = request.json
    format = data.get('format')

    if not format:
        return jsonify({'error': 'Format not provided'}), 400

    # Corrected file path
    filepath = r'C:\Rohit\Project\Time-Table\Backend\Download\entire_timetable.xlsx'
    
    try:
        # Load the data from the specified Excel file
        df = pd.read_excel(filepath)

        # Ensure the dataframe has the required columns
        required_columns = ['Teacher', 'Subject', 'Department', 'Class', 'Day', 'Lecture Time']
        if not all(col in df.columns for col in required_columns):
            return jsonify({'error': 'Missing required columns in timetable data'}), 400

        # Generate the file in the requested format
        if format == 'csv':
            return generate_csv(df)
        elif format == 'word':
            return generate_word(df)
        elif format == 'pdf':
            return generate_pdf(df)
        else:
            return jsonify({'error': 'Unsupported format'}), 400
    except Exception as e:
        return jsonify({'error': f'Error loading the file: {str(e)}'}), 500

# Function to generate the timetable as a CSV file
def generate_csv(df):
    try:
        # Ensure the required columns are in the dataframe
        df = df[['Teacher', 'Subject', 'Day', 'Lecture Time', 'Department', 'Class']]
        
        # Use BytesIO to create an in-memory binary file-like object
        csv_output = io.BytesIO()
        
        # Write the dataframe to this in-memory binary object as CSV
        df.to_csv(csv_output, index=False)
        
        # Rewind to the start of the in-memory binary file
        csv_output.seek(0)
        
        # Send the CSV file to the user as an attachment
        return send_file(csv_output, as_attachment=True, download_name="entire_timetable.csv", mimetype="text/csv")
    except Exception as e:
        print(f"Error generating CSV: {str(e)}")  # Log any CSV generation errors
        raise e

if __name__ == '__main__':
    app.run(debug=True)
