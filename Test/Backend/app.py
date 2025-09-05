from flask import Flask, jsonify, request, send_file
from flask_cors import CORS
import pyodbc
import pandas as pd
from io import BytesIO
from docx import Document  # For Word document generation
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib import colors

app = Flask(__name__)
CORS(app)

# Database connection
def get_db_connection():
    try:
        connection = pyodbc.connect(
            r'DRIVER={ODBC Driver 17 for SQL Server};'
            r'SERVER=LAPTOP-38HJSU5G\SQLEXPRESS;'
            r'DATABASE=college;'
            r'Trusted_Connection=yes;'
        )
        return connection
    except pyodbc.Error as e:
        print(f"Error connecting to database: {e}")
        return None

# Truncate and upload new teacher data
@app.route('/upload', methods=['POST'])
def upload_teacher_data():
    file = request.files['file']
    if not file:
        return jsonify({"error": "No file uploaded"}), 400

    try:
        df = pd.read_excel(file)  # Read the Excel file into a DataFrame
        connection = get_db_connection()
        if connection:
            cursor = connection.cursor()
            cursor.execute("TRUNCATE TABLE teacher_data")  # Clear the existing table
            connection.commit()

            for _, row in df.iterrows():
                cursor.execute("""
                    INSERT INTO teacher_data (ID, Teacher, Subjects, Class, Department, Lecture, Practical)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                """, row['ID'], row['Teacher'], row['Subjects'], row['Class'], row['Department'], row['Lecture'], row['Practical'])
            
            connection.commit()
            connection.close()
            return jsonify({"message": "Data uploaded successfully!"})
        else:
            return jsonify({"error": "Failed to connect to database"}), 500
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"error": "Failed to process file"}), 500

@app.route('/generate', methods=['POST'])
def generate_timetable():
    try:
        connection = get_db_connection()
        if not connection:
            return jsonify({"error": "Failed to connect to database"}), 500

        cursor = connection.cursor()

        # Fetch all teachers and their subjects
        cursor.execute("SELECT * FROM teacher_data")
        teachers = cursor.fetchall()

        # Fetch all unique class and department combinations
        cursor.execute("SELECT DISTINCT Class, Department FROM teacher_data")
        classes_departments = cursor.fetchall()

        # Time slots and days
        time_slots = [
            "9:00-10:00", "10:00-11:00", "11:00-12:00", "12:00-1:00",
            "1:15-2:15", "2:15-3:15", "3:15-4:15"
        ]
        days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday"]

        # Initialize tracking structures
        occupied_slots = {}  # Tracks time slots already assigned to each teacher
        schedule = []        # Final schedule to be inserted into the database

        for class_department in classes_departments:
            class_name, department = class_department

            for day in days:
                # Reset daily load for each day
                daily_load = {teacher.Teacher: {"lecture": 0, "practical": 0} for teacher in teachers}

                for time in time_slots:
                    for teacher in teachers:
                        teacher_name = teacher.Teacher
                        time_slot = f"{day} {time}"

                        # Check if the teacher is already occupied in this time slot
                        if occupied_slots.get(teacher_name, set()).intersection({time_slot}):
                            continue

                        # Assign lectures and practicals if limits are not exceeded
                        if teacher_name not in occupied_slots:
                            occupied_slots[teacher_name] = set()

                        if daily_load[teacher_name]["lecture"] < 2:
                            # Schedule a lecture
                            schedule.append({
                                "Teacher": teacher_name,
                                "Subjects": teacher.Subjects,
                                "Class": class_name,
                                "Department": department,
                                "D_name": day,
                                "Time_Slot": time,
                                "Lecture": "Yes",
                                "Practical": "No"
                            })
                            daily_load[teacher_name]["lecture"] += 1
                            occupied_slots[teacher_name].add(time_slot)

                        elif daily_load[teacher_name]["practical"] < 1:
                            # Schedule a practical
                            schedule.append({
                                "Teacher": teacher_name,
                                "Subjects": teacher.Subjects,
                                "Class": class_name,
                                "Department": department,
                                "D_name": day,
                                "Time_Slot": time,
                                "Lecture": "No",
                                "Practical": "Yes"
                            })
                            daily_load[teacher_name]["practical"] += 1
                            occupied_slots[teacher_name].add(time_slot)

        # Insert the updated schedule into the database
        cursor.execute("TRUNCATE TABLE schedule")
        for entry in schedule:
            cursor.execute("""
                INSERT INTO schedule (Teacher, Subjects, Class, Department, D_name, Time_Slot, Lecture, Practical)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """, entry['Teacher'], entry['Subjects'], entry['Class'],
               entry['Department'], entry['D_name'], entry['Time_Slot'],
               entry['Lecture'], entry['Practical'])

        connection.commit()
        connection.close()

        return jsonify({"message": "Schedule generated successfully!"})
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"error": "Failed to generate schedule"}), 500


# Serve dropdown options for Department, Class, Teacher
@app.route('/dropdown/department', methods=['GET'])
def get_department_dropdown():
    return get_dropdown_options('Department')

@app.route('/dropdown/class', methods=['GET'])
def get_class_dropdown():
    return get_dropdown_options('Class')

@app.route('/dropdown/teacher', methods=['GET'])
def get_teacher_dropdown():
    return get_dropdown_options('Teacher')

# Generalized function to fetch distinct values
def get_dropdown_options(field):
    try:
        connection = get_db_connection()
        if not connection:
            return jsonify({"error": "Failed to connect to database"}), 500

        cursor = connection.cursor()
        cursor.execute(f"SELECT DISTINCT {field} FROM teacher_data ORDER BY {field} ASC")
        options = [row[0] for row in cursor.fetchall()]
        connection.close()
        return jsonify({"options": options})
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"error": "Failed to fetch dropdown data"}), 500

# Download timetable in table format (CSV, PDF, DOCX)
@app.route('/download', methods=['GET'])
def download_timetable():
    file_type = request.args.get('type', 'csv').lower()  # Ensure the type is lowercase
    try:
        connection = get_db_connection()
        if not connection:
            return jsonify({"error": "Failed to connect to database"}), 500

        # Fetch timetable from database
        df = pd.read_sql("SELECT * FROM schedule", connection)
        connection.close()

        # For CSV download
        if file_type == 'csv':
            output = BytesIO()
            df.to_csv(output, index=False)
            output.seek(0)
            return send_file(output, as_attachment=True, download_name="timetable.csv", mimetype="text/csv")

        # For PDF download
        elif file_type == 'pdf':
            output = BytesIO()
            doc = SimpleDocTemplate(output, pagesize=letter)
            elements = []

            # Prepare data for the table
            data = [df.columns.tolist()]  # Header row
            data.extend(df.values.tolist())  # Add table rows

            # Create the table with styles
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('BACKGROUND', (0, 1), (-1, -1), colors.whitesmoke)
            ]))

            elements.append(table)
            doc.build(elements)

            output.seek(0)
            return send_file(output, as_attachment=True, download_name="timetable.pdf", mimetype="application/pdf")

        # For Word document download
        elif file_type == 'docx':
            doc = Document()

            doc.add_heading('Timetable', 0)

            # Add table and header
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, column_name in enumerate(df.columns):
                hdr_cells[i].text = column_name

            # Add rows and apply borders
            for index, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

            # Save to BytesIO
            output = BytesIO()
            doc.save(output)
            output.seek(0)

            return send_file(output, as_attachment=True, download_name="timetable.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        return jsonify({"error": "Invalid file type"}), 400
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"error": "Failed to download timetable"}), 500

if __name__ == '__main__':
    app.run(debug=True)
